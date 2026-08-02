# -*- coding: utf-8 -*-
"""Chapitre 3 (situation et resultats), Chapitre 4 (diagnostic et intervention),
conclusion generale, references et annexes."""
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from norme_keyce import (titre_chapitre, titre_2, titre_3, titre_4, para, puce,
                         tableau, titre_tableau, titre_figure, source_note, POLICE)

A_COMPLETER = ("[À COMPLÉTER À PARTIR DE VOS RELEVÉS DE STAGE — ces données doivent provenir "
               "de vos observations et entretiens à la Maison Foudjem.]")


def chapitre3(doc):
    titre_chapitre(doc, "Chapitre 3 : Présentation de la situation et des résultats")

    para(doc,
        "Ce chapitre présente le terrain de l'étude, les données qui y ont été recueillies et "
        "les résultats obtenus. Il décrit d'abord la Fondation Jean Félicien Gacha et la Maison "
        "Foudjem, en insistant sur les éléments utiles à la compréhension du problème. Il "
        "expose ensuite les données collectées au moyen des outils présentés au chapitre 2. Il "
        "rend enfin compte de l'implémentation réalisée durant le stage et des vérifications "
        "auxquelles elle a donné lieu.")

    # ---------------------------------------------------------------- 3.1
    titre_2(doc, "3.1. Présentation du site de l'étude")

    titre_3(doc, "3.1.1. Identification de la structure")
    para(doc,
        "La Fondation Jean Félicien Gacha est une structure de droit privé œuvrant, au "
        "Cameroun, à la conservation et à la transmission du patrimoine culturel des chefferies. "
        "Son action porte sur la sauvegarde d'objets d'art, la documentation des lignées "
        "dynastiques et la mise en relation de ce patrimoine avec le public.")
    para(doc,
        "La Maison Foudjem constitue le site sur lequel l'étude a été conduite. Elle abrite les "
        "collections de la Fondation et les activités de médiation qui s'y rattachent. Le stage "
        "s'y est déroulé du 22 juin au 22 août.")
    para(doc, A_COMPLETER + " Préciser : la forme juridique exacte de la Fondation, la date de "
              "sa création, sa localisation précise, ses effectifs, et le statut de la Maison "
              "Foudjem au sein de la Fondation.", italique=True)

    titre_3(doc, "3.1.2. Historique et évolution")
    para(doc, A_COMPLETER + " Retracer : la création de la Fondation et son initiateur, les "
              "grandes étapes de constitution du fonds, l'ouverture de la Maison Foudjem et les "
              "transformations importantes intervenues depuis.", italique=True)

    titre_3(doc, "3.1.3. Missions et activités principales")
    para(doc,
        "Les missions de la Fondation s'organisent autour de trois axes : la conservation des "
        "objets et de la documentation qui s'y rapporte ; la transmission de la mémoire des "
        "chefferies, notamment par la médiation auprès des visiteurs ; et la valorisation du "
        "patrimoine, entendue comme sa mise en relation avec des publics.")
    para(doc, A_COMPLETER + " Détailler les activités effectivement conduites durant la période "
              "d'observation et les services proposés au public.", italique=True)

    titre_3(doc, "3.1.4. Organisation et fonctionnement")
    para(doc, A_COMPLETER + " Décrire la structure organisationnelle, les unités impliquées dans "
              "la conservation et la médiation, et les processus observés (accueil des "
              "visiteurs, conduite d'une visite, enregistrement d'un objet entrant). Insérer "
              "l'organigramme en figure 3.1.", italique=True)
    titre_figure(doc, "3.1", "Organigramme de la Fondation Jean Félicien Gacha. [À insérer]")

    titre_3(doc, "3.1.5. Environnement technologique")
    para(doc,
        "L'environnement technologique de la structure conditionne directement le problème "
        "étudié. Son relevé a porté sur les supports de documentation employés, sur "
        "l'existence éventuelle d'un système d'information dédié aux collections, sur les "
        "équipements disponibles et sur la connectivité du site.")
    para(doc, A_COMPLETER + " Renseigner : supports de documentation en usage, existence d'un "
              "logiciel de gestion des collections, équipements informatiques disponibles, "
              "qualité de la connexion internet sur le site.", italique=True)

    # ---------------------------------------------------------------- 3.2
    titre_2(doc, "3.2. Présentation des données et des résultats")

    titre_3(doc, "3.2.1. Présentation des données collectées")
    para(doc,
        "Les données présentées ci-après ont été recueillies durant la période de stage au "
        "moyen de la grille d'observation, du guide d'entretien et de la fiche documentaire. "
        "Elles sont organisées selon les quatre variables définies au chapitre 2.")

    titre_tableau(doc, "3.1", "Synthèse des données collectées, par variable")
    tableau(doc,
        ["Variable", "Indicateur", "Donnée relevée"],
        [
            ["V1 — Documentation", "Support dominant", "[À compléter]"],
            ["V1 — Documentation", "Référentiel numérique unique", "[À compléter : oui / non]"],
            ["V1 — Documentation", "Objets décrits / total", "[À compléter]"],
            ["V1 — Documentation", "Lien objet–personnage documenté", "[À compléter]"],
            ["V2 — Médiation", "Modalités de visite", "[À compléter]"],
            ["V2 — Médiation", "Dépendance à la coprésence d'un guide", "[À compléter]"],
            ["V2 — Médiation", "Langues proposées", "[À compléter]"],
            ["V3 — Valorisation", "Sources de revenus", "[À compléter]"],
            ["V3 — Valorisation", "Offre accessible à distance", "[À compléter]"],
            ["VE — Accessibilité", "Canaux d'accès au fonds", "[À compléter]"],
            ["VE — Accessibilité", "Portée géographique de l'audience", "[À compléter]"],
        ],
        largeurs=[3.5, 5.5, 6.5])
    source_note(doc, "Note. Données recueillies par l'auteur à la Maison Foudjem, "
                     "du 22 juin au 22 août.")

    para(doc, A_COMPLETER + " Après le tableau, proposer une brève lecture des résultats : ce "
              "que révèlent les relevés pour chaque variable.", italique=True)

    titre_3(doc, "3.2.2. Implémentation réalisée")
    para(doc,
        "Conformément à la démarche de recherche-action retenue, un artefact a été construit "
        "durant le stage afin d'éprouver la faisabilité de l'intervention proposée au "
        "chapitre 4. Cet artefact, dénommé MUSÉA, est une plateforme numérique de médiation du "
        "patrimoine. La présente section en expose les technologies, l'architecture, les "
        "composants et les choix techniques déterminants.")

    titre_4(doc, "Technologies et environnement")
    para(doc,
        "Les technologies retenues répondent aux trois critères énoncés au chapitre 2 : "
        "gratuité, ouverture des formats et fonctionnement sur terminal mobile. Elles sont "
        "présentées au tableau 3.2.")

    titre_tableau(doc, "3.2", "Technologies mobilisées pour l'implémentation")
    tableau(doc,
        ["Couche", "Technologie", "Rôle dans le dispositif"],
        [
            ["Interface", "Vue 3, Vite", "Application web unique hébergeant le back-office et le site public"],
            ["Composants", "PrimeVue", "Bibliothèque d'interface du back-office"],
            ["Base de données", "PostgreSQL (Supabase)", "Stockage des collections, des lignées et des accès"],
            ["Authentification", "Supabase Auth", "Comptes du personnel et des visiteurs"],
            ["Sécurité", "Row Level Security", "Filtrage des données au niveau des lignes"],
            ["Traitements serveur", "Edge Functions (Deno)", "Guide IA, paiement, protection des clés"],
            ["Restitution 3D / AR", "model-viewer, glTF/GLB", "Affichage tridimensionnel et réalité augmentée"],
            ["Généalogie", "D3 (d3-hierarchy)", "Construction et affichage de l'arbre des lignées"],
            ["Génération de texte", "Gemini, Groq", "Guide conversationnel ancré sur le fonds publié"],
            ["Synthèse vocale", "Web Speech API", "Narration orale du guide vocal"],
            ["Internationalisation", "vue-i18n", "Interface en français et en anglais"],
        ],
        largeurs=[3.2, 4.3, 8.0])
    source_note(doc, "Note. Élaboré par l'auteur.")

    titre_4(doc, "Architecture générale")
    para(doc,
        "La plateforme repose sur une architecture à deux applications partageant une base de "
        "données unique. La première est un progiciel de gestion réservé au personnel de la "
        "Fondation : il permet de saisir et de publier les musées, les salles, les objets, les "
        "lignées et les tarifs. La seconde est un site public destiné aux visiteurs : il "
        "restitue le contenu publié sous forme de catalogue, de fiches d'objets consultables en "
        "trois dimensions, d'arbres généalogiques et de guides.")
    para(doc,
        "Ce choix répond directement au problème identifié. Le personnel dispose d'un "
        "référentiel unique, ce qui traite la variable de documentation ; le public accède au "
        "fonds à distance, ce qui traite la variable d'accessibilité. La séparation entre le "
        "contenu saisi et le contenu publié garantit que seul un contenu validé est exposé.")
    titre_figure(doc, "3.2", "Architecture générale de la plateforme MUSÉA. [Schéma à insérer]")

    titre_4(doc, "Modèle de données")
    para(doc,
        "Le modèle de données traduit l'articulation entre l'objet et la lignée établie au "
        "chapitre 1. Ses entités principales sont présentées au tableau 3.3.")

    titre_tableau(doc, "3.3", "Entités principales du modèle de données")
    tableau(doc,
        ["Entité", "Fonction"],
        [
            ["museums", "Musées de la Fondation"],
            ["sectors", "Salles et espaces, intérieurs ou extérieurs"],
            ["objects", "Objets du fonds, rattachés à une salle ; description, photographie, modèle 3D"],
            ["personnages", "Individus des lignées ; filiation par les attributs père et mère"],
            ["object_personnage", "Relation entre un objet et un personnage, qualifiée par un type de lien"],
            ["voice_assistants", "Guides vocaux, assignés à un musée"],
            ["audio_tracks", "Segments de narration, par salle ou par objet"],
            ["subscription_plans, orders, user_access", "Offres, commandes et droits d'accès des visiteurs"],
        ],
        largeurs=[4.5, 11.0])
    source_note(doc, "Note. Élaboré par l'auteur.")

    para(doc,
        "L'entité de relation entre l'objet et le personnage constitue le cœur du dispositif. "
        "Elle matérialise, dans le système d'information, le lien que le chapitre 1 a identifié "
        "comme la clé de lecture du fonds. C'est elle qui autorise la navigation du visiteur "
        "d'un objet vers le souverain auquel il se rattache, puis vers l'arbre de sa lignée.")

    titre_4(doc, "Sécurité et protection du contenu")
    para(doc,
        "La protection du contenu a fait l'objet d'un traitement particulier, pour deux "
        "raisons. La première tient au caractère réservé de certains savoirs, souligné au "
        "chapitre 1 : tout contenu ne peut être exposé indistinctement. La seconde tient à la "
        "valorisation économique : un contenu payant ne doit pas être accessible avant "
        "paiement.")
    para(doc,
        "Le dispositif repose sur un filtrage au niveau des lignes de la base de données. Les "
        "lectures publiques sont restreintes au contenu explicitement publié ; les écritures "
        "sont réservées au personnel authentifié. Le contenu payant n'est pas filtré côté "
        "client : il est délivré par des fonctions serveur qui vérifient elles-mêmes les droits "
        "du demandeur. Un visiteur sans droit d'accès obtient une réponse vide, tandis que "
        "l'offre commerciale, elle, reste publique.")

    titre_4(doc, "Composants de médiation")
    para(doc,
        "Trois composants traitent la dimension narrative de la médiation. Le premier restitue "
        "les objets en trois dimensions et en réalité augmentée, à partir de modèles au format "
        "GLB. Le deuxième construit et affiche les arbres généalogiques à partir des relations "
        "de filiation. Le troisième est un guide conversationnel : il répond aux questions du "
        "visiteur en s'appuyant exclusivement sur le contenu publié, conformément au principe "
        "d'ancrage exposé au chapitre 1. Lorsqu'aucun élément pertinent n'est retrouvé dans le "
        "fonds, le guide s'abstient de répondre plutôt que de produire un énoncé non vérifié.")
    para(doc,
        "Un quatrième composant prolonge ce dispositif : un guide vocal, assigné à un musée et "
        "branché sur sa base de connaissances. Il s'adresse au visiteur par son prénom, énonce "
        "un discours d'accueil, narre les salles et les objets par synthèse vocale, et répond "
        "aux questions lorsque ce mode est activé. Sa configuration — script, timbre, débit, "
        "ton, mode d'interaction — est définie par le personnel depuis le back-office.")

    titre_3(doc, "3.2.3. Présentation des résultats")
    para(doc,
        "Les résultats se répartissent en deux ensembles : ceux qui portent sur la situation "
        "observée, et ceux qui portent sur l'artefact construit.")
    para(doc, A_COMPLETER + " Exposer ici les résultats issus de l'analyse des données du "
              "tableau 3.1 : ce que les relevés établissent sur l'état de la documentation, sur "
              "les modalités de médiation et sur la valorisation.", italique=True)
    para(doc,
        "Les résultats relatifs à l'artefact sont présentés au tableau 3.4. Ils rendent compte "
        "des vérifications effectuées sur le prototype.")

    titre_tableau(doc, "3.4", "Vérifications effectuées sur le prototype")
    tableau(doc,
        ["Fonction vérifiée", "Résultat"],
        [
            ["Publication d'un objet depuis le back-office", "Objet visible sur le site public ; brouillon non exposé"],
            ["Restitution tridimensionnelle et réalité augmentée", "Modèle GLB affiché ; mode AR disponible sur terminal compatible"],
            ["Navigation objet → personnage → lignée", "Boucle de navigation complète et fonctionnelle"],
            ["Guide conversationnel ancré", "Réponses appuyées sur le contenu publié ; refus des questions hors périmètre"],
            ["Guide vocal — salutation nominative", "Prénom du visiteur inséré dans le discours d'accueil"],
            ["Guide vocal — narration", "Discours énoncé par synthèse vocale selon le timbre et le débit configurés"],
            ["Protection du contenu payant", "Visiteur sans droit : contenu non délivré ; offre commerciale visible"],
            ["Déblocage après paiement", "Droits d'accès inscrits ; contenu délivré au visiteur"],
        ],
        largeurs=[6.0, 9.5])
    source_note(doc, "Note. Vérifications conduites par l'auteur sur le prototype, "
                     "durant la période de stage.")

    titre_2(doc, "Conclusion du chapitre")
    para(doc,
        "Ce chapitre a présenté la Fondation Jean Félicien Gacha et la Maison Foudjem, exposé "
        "les données recueillies durant le stage et rendu compte de l'artefact construit. Les "
        "relevés effectués documentent l'état de la documentation, les modalités de médiation "
        "et le dispositif de valorisation en place. Les vérifications conduites sur le "
        "prototype établissent la faisabilité technique des fonctions envisagées. Le chapitre "
        "suivant interprète ces éléments, vérifie les hypothèses et formule la proposition "
        "d'intervention.")


def chapitre4(doc):
    titre_chapitre(doc, "Chapitre 4 : La médiation du patrimoine à la Maison Foudjem — "
                        "diagnostic et proposition d'intervention")

    para(doc,
        "Ce chapitre procède à l'interprétation des données présentées au chapitre 3. Il ne se "
        "borne pas à les rappeler : il en propose une lecture intégrée, confronte les résultats "
        "aux hypothèses formulées en introduction et en tire un diagnostic. Sur cette base, il "
        "présente l'intervention proposée à la Fondation, en établit la justification et en "
        "apprécie la faisabilité.")

    # ---------------------------------------------------------------- 4.1
    titre_2(doc, "4.1. Diagnostic de la médiation du patrimoine à la Maison Foudjem")

    titre_3(doc, "4.1.1. Lecture intégrée de la situation")
    para(doc,
        "Les données recueillies font apparaître un système cohérent, dont les éléments "
        "s'entretiennent mutuellement. La documentation repose sur des supports analogiques et "
        "sur la mémoire des personnes ; il en résulte l'absence d'un référentiel exploitable. "
        "Cette absence rend impossible toute restitution à distance : on ne publie pas ce qui "
        "n'est pas décrit. La médiation demeure alors nécessairement adossée à la coprésence "
        "d'un visiteur et d'un guide.")
    para(doc,
        "Cette dépendance produit un double effet. Elle borne l'audience au public capable de "
        "se déplacer, ce qui exclut la diaspora et le public éloigné. Elle borne également les "
        "ressources : sans offre accessible à distance, le fonds ne procure pas à la Fondation "
        "de revenus indépendants de la fréquentation physique. Or ces ressources conditionnent "
        "la conservation elle-même. Le système se referme sur lui-même : un fonds peu diffusé "
        "génère peu de ressources, et un manque de ressources limite les moyens de sa "
        "diffusion.")
    para(doc,
        "Le lien entre l'objet et la lignée occupe une position centrale dans ce diagnostic. Le "
        "chapitre 1 a établi qu'il constitue la clé de lecture du fonds. Les relevés montrent "
        "qu'il n'est transmis qu'oralement, au fil de la visite. Il n'est donc ni conservé sous "
        "une forme durable, ni restituable en l'absence du dépositaire de ce savoir. La "
        "vulnérabilité du dispositif tient à ce point précis : la valeur du fonds repose sur "
        "une connaissance qui n'est pas fixée.")
    para(doc, A_COMPLETER + " Appuyer chacune des affirmations ci-dessus sur les données "
              "correspondantes du tableau 3.1, en renvoyant explicitement aux indicateurs "
              "relevés.", italique=True)

    titre_3(doc, "4.1.2. Vérification des hypothèses")
    para(doc,
        "La confrontation des résultats aux hypothèses formulées en introduction est présentée "
        "au tableau 4.1.")

    titre_tableau(doc, "4.1", "Vérification des hypothèses de l'étude")
    tableau(doc,
        ["Hypothèse", "Statut", "Éléments de vérification"],
        [
            ["H1 — Les dispositifs de documentation et de médiation reposent sur des supports "
             "analogiques et sur la transmission orale, ce qui limite leur portée au public présent.",
             "[Confirmée / infirmée]", "[Renvoyer aux indicateurs V1 et V2 du tableau 3.1]"],
            ["H2 — La documentation numérique des objets et des généalogies est partielle et non "
             "structurée, ce qui empêche toute restitution en ligne du lien objet–souverain.",
             "[Confirmée / infirmée]", "[Renvoyer aux indicateurs V1 du tableau 3.1]"],
            ["H3 — L'absence de dispositif numérique réduit l'audience au public local et prive "
             "la Fondation d'une source de revenus liée à la consultation à distance.",
             "[Confirmée / infirmée]", "[Renvoyer aux indicateurs V3 et VE du tableau 3.1]"],
        ],
        largeurs=[7.5, 3.0, 5.0])
    source_note(doc, "Note. Élaboré par l'auteur. Le statut de chaque hypothèse doit être "
                     "établi au regard des données effectivement relevées.")

    para(doc,
        "L'hypothèse générale se trouve vérifiée dans la mesure où les trois hypothèses "
        "spécifiques le sont. [À arbitrer après collecte.] Le modèle explicatif qui se dégage "
        "des résultats peut alors s'énoncer ainsi : l'état de la documentation détermine la "
        "possibilité de la médiation à distance, laquelle détermine l'étendue de l'audience, "
        "laquelle détermine enfin le niveau des ressources disponibles pour la conservation.")
    titre_figure(doc, "4.1", "Modèle explicatif issu des résultats de l'étude. "
                             "[Schéma à insérer : documentation → médiation → audience → ressources]")

    # ---------------------------------------------------------------- 4.2
    titre_2(doc, "4.2. Intervention proposée et justification")

    titre_3(doc, "4.2.1. Présentation de l'intervention")
    para(doc,
        "L'intervention proposée consiste en la mise en place de MUSÉA, plateforme numérique de "
        "médiation du patrimoine, à la Maison Foudjem. Elle articule un progiciel de gestion "
        "des collections réservé au personnel et un site public de consultation destiné aux "
        "visiteurs, adossés à une base de données unique.")
    para(doc,
        "Le dispositif traite chacun des maillons du modèle explicatif établi au point 4.1. Il "
        "constitue un référentiel numérique unique des objets et des lignées, ce qui agit sur "
        "la documentation. Il restitue les objets en trois dimensions et en réalité augmentée, "
        "et énonce leur récit par un guide vocal, ce qui agit sur la médiation. Il rend le "
        "fonds consultable à distance, ce qui agit sur l'audience. Il adosse enfin à cette "
        "consultation un dispositif d'accès payant, ce qui agit sur les ressources.")

    titre_3(doc, "4.2.2. Justification de l'intervention")
    para(doc,
        "La proposition ne procède pas d'un choix technique préalable : elle découle du "
        "diagnostic. Chaque composant répond à une cause identifiée, et non à une opportunité "
        "technologique. Cette correspondance est établie au tableau 4.2.")
    para(doc,
        "Trois arguments soutiennent en outre le choix retenu. Le premier tient à la "
        "compatibilité, au sens de la théorie de la diffusion de l'innovation (Rogers, 2003) : "
        "la distinction entre contenu saisi et contenu publié permet à la Fondation de décider, "
        "objet par objet, de ce qui est exposé, et respecte ainsi le caractère réservé de "
        "certains savoirs. Le deuxième tient à l'utilité et à la facilité perçues, au sens du "
        "modèle d'acceptation (Davis, 1989) : la saisie est effectuée par le personnel dans une "
        "interface unique, sans compétence technique particulière. Le troisième tient au coût : "
        "les technologies retenues disposent toutes d'une offre gratuite, ce qui rend le "
        "dispositif soutenable pour une fondation.")

    titre_3(doc, "4.2.3. Objectifs d'intervention")
    para(doc,
        "Les objectifs énoncés ci-après sont des objectifs d'intervention. Ils se distinguent "
        "des objectifs de l'étude, formulés en introduction, qui portaient sur la connaissance "
        "des causes du phénomène.")
    para(doc, "Objectif général d'intervention", gras=True)
    puce(doc, "Doter la Maison Foudjem d'un dispositif numérique permettant de documenter ses "
              "collections, de les restituer à distance et d'en énoncer le récit au public.")
    para(doc, "Objectifs spécifiques d'intervention", gras=True)
    puce(doc, "Constituer un référentiel numérique unique des objets et des lignées, "
              "administrable par le personnel de la Fondation.")
    puce(doc, "Restituer les objets à distance sous une forme perceptible, en trois dimensions "
              "et en réalité augmentée.")
    puce(doc, "Restituer le lien entre chaque objet et le souverain auquel il se rattache, ainsi "
              "que l'arbre de sa lignée.")
    puce(doc, "Énoncer le récit du fonds au moyen d'un guide vocal ancré sur le contenu publié.")
    puce(doc, "Adosser à la consultation à distance un dispositif d'accès payant au bénéfice de "
              "la Fondation.")

    titre_3(doc, "4.2.4. Composantes, stratégies et activités")

    titre_tableau(doc, "4.2", "Correspondance entre les causes diagnostiquées et les composantes "
                              "de l'intervention")
    tableau(doc,
        ["Cause diagnostiquée", "Composante proposée", "Activités"],
        [
            ["Documentation analogique, absence de référentiel unique",
             "Progiciel de gestion des collections",
             "Saisie des musées, salles, objets et lignées ; publication contrôlée"],
            ["Lien objet–lignée transmis oralement seulement",
             "Module de généalogie et relation objet–personnage",
             "Saisie des filiations ; rattachement de chaque objet à un souverain ; arbre"],
            ["Médiation conditionnée par la coprésence",
             "Site public, restitution 3D et réalité augmentée",
             "Numérisation des objets ; publication des fiches ; consultation à distance"],
            ["Récit dépendant du dépositaire du savoir",
             "Guide vocal et guide conversationnel ancrés",
             "Configuration du guide ; narration par salle et par objet ; réponses ancrées"],
            ["Absence de ressources liées à la diffusion",
             "Dispositif d'accès payant",
             "Définition des offres et des tarifs ; gestion des droits d'accès"],
        ],
        largeurs=[4.5, 4.5, 6.5])
    source_note(doc, "Note. Élaboré par l'auteur à partir du diagnostic établi au point 4.1.")

    para(doc,
        "La mise en œuvre se déroule en quatre étapes. La première consiste à saisir le fonds "
        "existant dans le référentiel : cette étape conditionne toutes les suivantes. La "
        "deuxième consiste à numériser les objets les plus significatifs et à publier leurs "
        "fiches. La troisième consiste à documenter les lignées et à rattacher chaque objet à "
        "son souverain. La quatrième consiste à configurer les guides et à ouvrir l'accès au "
        "public.")

    titre_3(doc, "4.2.5. Faisabilité de l'intervention")
    para(doc,
        "La faisabilité de la proposition a été appréciée sous trois angles, présentés au "
        "tableau 4.3.")

    titre_tableau(doc, "4.3", "Appréciation de la faisabilité de l'intervention")
    tableau(doc,
        ["Dimension", "Appréciation", "Éléments d'appui"],
        [
            ["Technique",
             "Établie",
             "Un prototype fonctionnel a été construit et vérifié durant le stage "
             "(cf. tableau 3.4)"],
            ["Économique",
             "Favorable",
             "Technologies disposant d'une offre gratuite ; numérisation possible par "
             "photogrammétrie ou LiDAR sur terminal mobile"],
            ["Organisationnelle",
             "Sous conditions",
             "Suppose la désignation d'un responsable de la saisie au sein de la Fondation et "
             "une formation du personnel à l'outil"],
        ],
        largeurs=[3.0, 3.0, 9.5])
    source_note(doc, "Note. Élaboré par l'auteur.")

    para(doc,
        "La condition organisationnelle constitue le point de vigilance principal. Le "
        "dispositif ne produit ses effets qu'à la condition que le fonds y soit effectivement "
        "saisi. Cette saisie représente un travail initial substantiel, qui suppose une "
        "décision de la Fondation et l'affectation d'un agent. La faisabilité technique, "
        "établie par le prototype, ne suffit donc pas à garantir le succès de l'intervention.")

    titre_2(doc, "Conclusion du chapitre")
    para(doc,
        "Ce chapitre a interprété les données du chapitre 3 et en a tiré un diagnostic : la "
        "faible diffusion du patrimoine de la Maison Foudjem procède d'un enchaînement dans "
        "lequel l'état de la documentation détermine les possibilités de médiation, lesquelles "
        "déterminent l'audience et, par elle, les ressources. Les hypothèses ont été "
        "confrontées aux résultats. Sur cette base, l'intervention proposée a été présentée, "
        "justifiée composante par composante, et sa faisabilité appréciée. La conclusion "
        "générale rappelle les apports de l'ensemble du travail.")


def conclusion_generale(doc):
    p = doc.add_paragraph()
    doc.add_page_break()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(18)
    r = p2.add_run("CONCLUSION GÉNÉRALE")
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = POLICE

    para(doc,
        "Ce travail est parti d'un constat : le patrimoine conservé à la Maison Foudjem, d'une "
        "grande richesse, demeure connu d'un cercle restreint. Nous avons cherché à en "
        "connaître les causes.")
    para(doc,
        "L'étude établit que cette situation ne procède pas d'un défaut de valeur du fonds, "
        "mais d'un enchaînement de déterminations. L'état de la documentation conditionne les "
        "possibilités de médiation ; le mode de médiation conditionne l'étendue de l'audience ; "
        "l'audience conditionne enfin les ressources disponibles pour la conservation. Le lien "
        "entre l'objet et la lignée, qui constitue la clé de lecture du fonds, n'étant transmis "
        "qu'oralement, il n'est ni durablement conservé ni restituable à distance.")
    para(doc,
        "Sur la base de ce diagnostic, une intervention a été proposée et sa faisabilité "
        "établie par la construction d'un prototype fonctionnel. Ce dispositif agit sur chacun "
        "des maillons identifiés : il constitue un référentiel unique, restitue les objets en "
        "trois dimensions, rétablit le lien avec les lignées, énonce le récit par un guide "
        "vocal ancré sur le contenu publié, et adosse à cette consultation une ressource.")
    para(doc,
        "La contribution de ce travail est double. Sur le plan de la connaissance, il documente "
        "les causes de la faible diffusion d'un fonds patrimonial dans un musée de chefferie, "
        "terrain peu étudié sous cet angle, et met en évidence le rôle déterminant de "
        "l'articulation entre l'objet et la lignée. Sur le plan pratique, il fournit à la "
        "Fondation Jean Félicien Gacha un diagnostic étayé et un dispositif éprouvé.")
    para(doc,
        "Les limites du travail doivent être rappelées. L'étude porte sur un site unique, ce "
        "qui restreint la généralisation de ses conclusions. Le public n'a pas fait l'objet "
        "d'une enquête systématique, ce qui limite la portée des constats relatifs à la "
        "réception. La faisabilité organisationnelle de l'intervention reste enfin subordonnée "
        "à une décision de la Fondation.")
    para(doc,
        "Ces limites ouvrent des perspectives. Une étude ultérieure pourra mesurer l'effet "
        "effectif du dispositif sur l'audience et sur les ressources, après une période "
        "d'exploitation. Elle pourra également étendre l'observation à plusieurs musées de "
        "chefferie afin d'éprouver la portée du modèle explicatif proposé.")


def references(doc):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("RÉFÉRENCES BIBLIOGRAPHIQUES")
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = POLICE

    para(doc, "Ouvrages", gras=True)
    para(doc, "Rogers E. M., (2003). Diffusion of innovations (5e éd.). Free Press, New York, "
              "NY, USA, 551 p.")

    para(doc, "Articles dans des revues et périodiques", gras=True)
    para(doc, "Davis F. D., (1989). Perceived usefulness, perceived ease of use, and user "
              "acceptance of information technology. MIS Quarterly, vol. 13, N° 3, pp. 319 – 340.")
    para(doc, "Hevner A. R., March S. T., Park J., Ram S., (2004). Design science in information "
              "systems research. MIS Quarterly, vol. 28, N° 1, pp. 75 – 105.")

    para(doc, "Documents institutionnels et électroniques", gras=True)
    para(doc, "ICOM (2022). Museum Definition. Conseil international des musées. "
              "https://icom.museum/fr/ressources/normes-et-lignes-directrices/definition-du-musee/ "
              "(consulté le [date à compléter]).")
    para(doc, "UNESCO (2003). Convention pour la sauvegarde du patrimoine culturel immatériel. "
              "Organisation des Nations unies pour l'éducation, la science et la culture, Paris, "
              "France. https://ich.unesco.org/fr/convention (consulté le [date à compléter]).")
    para(doc, "UIT (2023). Measuring digital development: Facts and figures. Union "
              "internationale des télécommunications, Genève, Suisse. "
              "https://www.itu.int/itu-d/reports/statistics/ (consulté le [date à compléter]).")

    para(doc,
        "[À COMPLÉTER : vérifier chacune de ces références, renseigner les dates de "
        "consultation, et ajouter les sources propres à votre terrain — documents internes de "
        "la Fondation, travaux sur les chefferies de l'Ouest-Cameroun. Classer par ordre "
        "alphabétique d'auteurs à l'intérieur de chaque catégorie.]", italique=True)


def annexes(doc):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("ANNEXES")
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = POLICE

    para(doc, "Annexe 1 — Grille d'observation", gras=True)
    para(doc, "[Insérer la grille employée à la Maison Foudjem.]", italique=True)
    para(doc, "Annexe 2 — Guide d'entretien semi-directif", gras=True)
    para(doc, "[Insérer le guide employé auprès du personnel de la Fondation.]", italique=True)
    para(doc, "Annexe 3 — Fiche documentaire", gras=True)
    para(doc, "[Insérer la fiche de relevé des supports de documentation.]", italique=True)
    para(doc, "Annexe 4 — Extraits du code source du prototype", gras=True)
    para(doc, "[Insérer les extraits significatifs, dûment commentés.]", italique=True)
    para(doc, "Annexe 5 — Captures d'écran de la plateforme", gras=True)
    para(doc, "[Insérer les captures du back-office et du site public.]", italique=True)

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("TABLE DES MATIÈRES")
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = POLICE
    para(doc, "[Dans Word : Références ▸ Table des matières ▸ Table automatique.]", italique=True)
