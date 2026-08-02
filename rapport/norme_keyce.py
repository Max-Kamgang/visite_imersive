# -*- coding: utf-8 -*-
"""
Mise en forme conforme au « GUIDE DE REDACTION SCIENTIFIQUE » KEYCE V3.0.

Normes appliquees (section II du guide) :
  - Page A4 ; marge gauche 3 cm, droite/haut/bas 2,5 cm
  - Police Times New Roman, taille 12
  - Interligne 1,5 ; ligne vide entre les paragraphes ; paragraphes justifies
  - Citations longues (> 40 mots) : retrait droite 1,4 cm, interligne simple
  - Pagination en bas a droite ; couverture non numerotee ;
    liminaires en chiffres romains ; corps en chiffres arabes (redemarre a 1)
  - Titres : chapitre MAJUSCULES GRAS 13 / sous-chapitre gras minuscule 13
             niveau 3 italique gras minuscule 12 / niveau 4 italique minuscule 12
  - Tableaux : titre AU-DESSUS ; Figures : titre EN DESSOUS (norme APA)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

POLICE = "Times New Roman"


# ---------------------------------------------------------------- bas niveau
def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    return e


def _champ(paragraphe, instruction):
    """Insere un champ Word (ex. PAGE) dans un paragraphe."""
    r1 = paragraphe.add_run()._r
    r1.append(_el("w:fldChar", **{"w:fldCharType": "begin"}))
    r2 = paragraphe.add_run()._r
    t = _el("w:instrText")
    t.set(qn("xml:space"), "preserve")
    t.text = instruction
    r2.append(t)
    r3 = paragraphe.add_run()._r
    r3.append(_el("w:fldChar", **{"w:fldCharType": "end"}))


def format_pagination(section, fmt=None, depart=None):
    """fmt : 'lowerRoman' | 'upperRoman' | 'decimal' ; depart : numero de depart."""
    sectPr = section._sectPr
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = _el("w:pgNumType")
        sectPr.append(pg)
    if fmt:
        pg.set(qn("w:fmt"), fmt)
    if depart is not None:
        pg.set(qn("w:start"), str(depart))


def numeroter_bas_droite(section):
    """Ajoute le numero de page en bas a droite du pied de page."""
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _champ(p, "PAGE")
    for r in p.runs:
        r.font.name = POLICE
        r.font.size = Pt(11)


def sans_pagination(section):
    section.footer.is_linked_to_previous = False
    if section.footer.paragraphs:
        section.footer.paragraphs[0].text = ""


# ---------------------------------------------------------------- document
def nouveau_document():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = POLICE
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), POLICE)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(12)          # ligne vide entre les paragraphes
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    s = doc.sections[0]
    _marges(s)
    return doc


def _marges(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)      # marge gauche 3 cm
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)


def nouvelle_section(doc):
    s = doc.add_section(WD_SECTION.NEW_PAGE)
    _marges(s)
    return s


# ---------------------------------------------------------------- titres
def titre_chapitre(doc, texte):
    """Chapitre : MAJUSCULES, GRAS, 13. Commence sur une nouvelle page."""
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(18)   # espace supplementaire apres le titre
    r = p.add_run(texte.upper())
    r.bold = True
    r.font.size = Pt(13)
    r.font.name = POLICE
    return p


def titre_2(doc, texte):
    """Sous-chapitre : gras, minuscule, 13. Espace supplementaire avant."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(texte)
    r.bold = True
    r.font.size = Pt(13)
    r.font.name = POLICE
    return p


def titre_3(doc, texte):
    """Niveau 3 : italique, gras, minuscule, 12."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(texte)
    r.bold = True
    r.italic = True
    r.font.size = Pt(12)
    r.font.name = POLICE
    return p


def titre_4(doc, texte):
    """Niveau 4 : italique, minuscule, 12."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(texte)
    r.italic = True
    r.font.size = Pt(12)
    r.font.name = POLICE
    return p


def titre_liminaire(doc, texte, saut=True):
    """Titre des pages liminaires (DEDICACE, REMERCIEMENTS...)."""
    if saut:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(texte.upper())
    r.bold = True
    r.font.size = Pt(13)
    r.font.name = POLICE
    return p


# ---------------------------------------------------------------- paragraphes
def para(doc, texte, gras=False, italique=False, centre=False, taille=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centre else WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(texte)
    r.bold = gras
    r.italic = italique
    r.font.size = Pt(taille)
    r.font.name = POLICE
    return p


def puce(doc, texte, niveau=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1 + niveau * 0.6)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(texte)
    r.font.size = Pt(12)
    r.font.name = POLICE
    return p


def citation_longue(doc, texte):
    """Citation > 40 mots : paragraphe a part, retrait droite 1,4, interligne simple."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(1.4)
    pf.right_indent = Cm(1.4)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(texte)
    r.font.size = Pt(12)
    r.font.name = POLICE
    return p


def titre_tableau(doc, numero, libelle):
    """Norme APA : le titre du TABLEAU se place AU-DESSUS."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"Tableau {numero}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = POLICE
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(libelle)
    r2.italic = True
    r2.font.size = Pt(11)
    r2.font.name = POLICE
    return p


def source_note(doc, texte):
    """Note / source sous un tableau ou une figure (norme APA)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(texte)
    r.italic = True
    r.font.size = Pt(10)
    r.font.name = POLICE
    return p


def titre_figure(doc, numero, libelle):
    """Norme APA : le titre de la FIGURE se place EN DESSOUS."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(f"Figure {numero}. ")
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = POLICE
    r2 = p.add_run(libelle)
    r2.italic = True
    r2.font.size = Pt(11)
    r2.font.name = POLICE
    return p


def tableau(doc, entetes, lignes, largeurs=None):
    t = doc.add_table(rows=1, cols=len(entetes))
    t.style = "Table Grid"
    for i, h in enumerate(entetes):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = POLICE
    for ligne in lignes:
        cells = t.add_row().cells
        for i, v in enumerate(ligne):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(v))
            r.font.size = Pt(11)
            r.font.name = POLICE
    if largeurs:
        for row in t.rows:
            for i, w in enumerate(largeurs):
                row.cells[i].width = Cm(w)
    return t
