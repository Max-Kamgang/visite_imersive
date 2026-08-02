# -*- coding: utf-8 -*-
"""
Genere le rapport de fin d'etudes aux normes KEYCE V3.0.
Theme : La mediation numerique du patrimoine des chefferies camerounaises —
        cas de la Maison Foudjem (Fondation Jean Felicien Gacha).
"""
import os
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from norme_keyce import (
    nouveau_document, nouvelle_section, format_pagination, numeroter_bas_droite,
    sans_pagination, titre_chapitre, titre_2, titre_3, titre_4, titre_liminaire,
    para, puce, tableau, titre_tableau, titre_figure, source_note, POLICE,
)

# ---------------------------------------------------------------- parametres
ETUDIANT = "DASSI KAMGANG Max Brian"
THEME = ("La médiation numérique du patrimoine des chefferies camerounaises : "
         "cas de la Maison Foudjem (Fondation Jean Félicien Gacha)")
OPTION = "Intelligence artificielle et big data"
ENCADREUR_ACAD = "Ing. Goudjou Blondon"
ENCADREUR_PRO = "Ing. Adrien Ghomsi"
ANNEE = "2025-2026"
VILLE = "Yaoundé, Cameroun"
STAGE = "du 22 juin au 22 août"

doc = nouveau_document()

# ================================================================ SECTION 1
# Couverture + page de garde : NON numerotees
s1 = doc.sections[0]
sans_pagination(s1)

# ---- PREMIERE DE COUVERTURE (modele 2.1 du guide)
t = tableau(doc, ["UNIVERSITE DE YAOUNDE I", "UNIVERSITY OF YAOUNDE I"], [
    ["ECOLE NATIONALE SUPERIEURE POLYTECHNIQUE DE YAOUNDE",
     "NATIONAL ADVANCED SCHOOL OF ENGINEERING OF YAOUNDE"],
    ["DEPARTEMENT DE GENIE INFORMATIQUE", "DEPARTMENT OF COMPUTER ENGINEERING"],
    ["INSTITUT SUPERIEUR KEYCE INFORMATIQUE & INTELLIGENCE ARTIFICIELLE",
     "HIGHER INSTITUTE KEYCE OF COMPUTER SCIENCE AND ARTIFICIAL INTELLIGENCE"],
], largeurs=[7.75, 7.75])
for row in t.rows:
    for c in row.cells:
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

para(doc, "")
para(doc, "[Insérer ici : logo ENSPY (gauche) — logo Université de Yaoundé I (droite)]",
     italique=True, centre=True, taille=10)
para(doc, "")

para(doc, THEME, gras=True, centre=True, taille=15)
para(doc, "")
para(doc, "Rapport de fin d'études", gras=True, centre=True, taille=13)
para(doc, "")
para(doc, "Présenté et soutenu par :", centre=True)
para(doc, ETUDIANT, gras=True, centre=True, taille=13)
para(doc, "")
para(doc, "En vue de l'obtention du", centre=True)
para(doc, "Diplôme de licence en génie informatique", gras=True, centre=True)
para(doc, f"Option : {OPTION}", centre=True)
para(doc, "")
para(doc, "Sous l'encadrement de :", centre=True)
para(doc, f"{ENCADREUR_ACAD}, Encadreur académique", gras=True, centre=True)
para(doc, f"{ENCADREUR_PRO}, Encadreur professionnel", gras=True, centre=True)
para(doc, "")
para(doc, "Devant le jury composé de :", centre=True)
para(doc, "Président : Prénom et Nom, Grade", centre=True, taille=11)
para(doc, "Rapporteur : Prénom et Nom, Grade", centre=True, taille=11)
para(doc, "Membres : Prénom et Nom, Grade", centre=True, taille=11)
para(doc, "Invité : Prénom et Nom, Titre ou fonction", centre=True, taille=11)
para(doc, "")
para(doc, f"Année académique : {ANNEE}", centre=True)
para(doc, "Date de la soutenance : ……………………", centre=True)

# ---- PAGE DE GARDE (modele 2.2 du guide)
doc.add_page_break()
para(doc, "")
para(doc, "[LOGO ENSPY]     [LOGO Université de Yaoundé I]     [LOGO KEYCE INFORMATIQUE]     "
          "[LOGO Fondation Jean Félicien Gacha]", italique=True, centre=True, taille=10)
para(doc, "")
para(doc, "")
para(doc, THEME, gras=True, centre=True, taille=15)
para(doc, "")
para(doc, "")
para(doc, "Rapport en vue de l'obtention du diplôme de :", centre=True)
para(doc, "Licence en génie informatique", gras=True, centre=True, taille=13)
para(doc, f"Option : {OPTION}", centre=True)
para(doc, "")
para(doc, "Présenté par :", centre=True)
para(doc, ETUDIANT, gras=True, centre=True, taille=13)
para(doc, "")
para(doc, "")
para(doc, "")
para(doc, VILLE, gras=True, centre=True)
para(doc, f"Année académique {ANNEE}", gras=True, centre=True)

# ================================================================ SECTION 2
# Liminaires : chiffres romains
s2 = nouvelle_section(doc)
format_pagination(s2, fmt="upperRoman", depart=1)
numeroter_bas_droite(s2)

# ---- DEDICACE (le guide : uniquement une personne ou un groupe, sans phrase)
titre_liminaire(doc, "Dédicace", saut=False)
para(doc, "")
para(doc, "")
para(doc, "À la famille Kamgang", italique=True, centre=True, taille=14)

# ---- REMERCIEMENTS (ordre du guide : jury/encadreurs, ecole, entreprise, autres)
titre_liminaire(doc, "Remerciements")
para(doc, "La réalisation de ce travail a bénéficié du concours de plusieurs personnes et "
          "institutions envers lesquelles nous exprimons notre reconnaissance.")
para(doc, "Nous remercions d'abord les membres du jury, qui ont accepté d'évaluer ce travail "
          "et dont les observations contribuent à en améliorer la qualité.")
para(doc, f"Notre gratitude va particulièrement à {ENCADREUR_ACAD}, notre encadreur académique, "
          "pour la rigueur méthodologique qu'il nous a transmise, sa disponibilité et ses "
          "orientations tout au long de ce travail.")
para(doc, f"Nous remercions également {ENCADREUR_PRO}, notre encadreur professionnel, pour son "
          "accompagnement technique au quotidien, la confiance qu'il nous a accordée et la "
          "qualité de son suivi durant toute la période de stage.")
para(doc, "Nos remerciements s'adressent ensuite à l'Institut Supérieur KEYCE Informatique & "
          "Intelligence Artificielle, ainsi qu'à l'École Nationale Supérieure Polytechnique de "
          "Yaoundé et à l'Université de Yaoundé I, pour la formation reçue et pour "
          "l'encadrement dont nous avons bénéficié.")
para(doc, "Nous exprimons notre reconnaissance à la Fondation Jean Félicien Gacha, et "
          "singulièrement à la Maison Foudjem, qui nous a ouvert ses portes, donné accès à ses "
          "collections et associé à ses réflexions sur la valorisation du patrimoine.")
para(doc, "Nous remercions enfin la famille Kamgang pour son soutien constant, ainsi que "
          "toutes les personnes qui, de près ou de loin, ont contribué à l'aboutissement de ce "
          "travail.")

# ---- LISTE DES SIGLES ET ABREVIATIONS (ordre alphabetique)
titre_liminaire(doc, "Liste des sigles et abréviations")
sigles = [
    ("API", "Application Programming Interface (interface de programmation applicative)"),
    ("AR", "Augmented Reality (réalité augmentée)"),
    ("CRUD", "Create, Read, Update, Delete (créer, lire, mettre à jour, supprimer)"),
    ("ENSPY", "École Nationale Supérieure Polytechnique de Yaoundé"),
    ("ERP", "Enterprise Resource Planning (progiciel de gestion intégré)"),
    ("FJFG", "Fondation Jean Félicien Gacha"),
    ("GLB / glTF", "Graphics Language Transmission Format (format d'échange de modèles 3D)"),
    ("ICOM", "International Council of Museums (Conseil international des musées)"),
    ("IA", "Intelligence artificielle"),
    ("KEYCE", "Institut Supérieur KEYCE Informatique & Intelligence Artificielle"),
    ("LiDAR", "Light Detection and Ranging (détection et télémétrie par la lumière)"),
    ("RLS", "Row Level Security (sécurité au niveau des lignes)"),
    ("SI", "Système d'information"),
    ("TAM", "Technology Acceptance Model (modèle d'acceptation de la technologie)"),
    ("TTS", "Text To Speech (synthèse vocale)"),
    ("UNESCO", "United Nations Educational, Scientific and Cultural Organization"),
    ("UY1", "Université de Yaoundé I"),
]
tableau(doc, ["Sigle", "Signification"], sigles, largeurs=[3.5, 12])

# ---- RESUME (300-400 mots + 5 mots cles)
titre_liminaire(doc, "Résumé")
para(doc,
    "Le patrimoine des chefferies de l'Ouest-Cameroun constitue un fonds culturel majeur : "
    "objets d'art, insignes du pouvoir, architectures et lignées dynastiques y forment un "
    "ensemble indissociable, où chaque objet tire son sens de l'histoire du souverain auquel "
    "il est rattaché. Conservé pour l'essentiel dans des musées de chefferie, ce patrimoine "
    "demeure toutefois peu accessible : sa documentation reste largement analogique, sa "
    "médiation dépend de la présence physique du visiteur et d'un guide, et le lien entre les "
    "objets et les généalogies qui les portent n'est que rarement restitué. La Maison Foudjem, "
    "musée de la Fondation Jean Félicien Gacha, illustre cette situation.")
para(doc,
    "Cette étude porte sur la médiation numérique de ce patrimoine. Elle cherche à comprendre "
    "pourquoi les collections de la Maison Foudjem restent faiblement accessibles et "
    "faiblement valorisées, malgré leur richesse. La question générale interroge les causes de "
    "cette situation ; les questions spécifiques portent sur les dispositifs actuels de "
    "documentation et de médiation, sur l'état de la numérisation des objets et des "
    "généalogies, et sur les effets de l'absence d'un dispositif numérique sur l'audience et "
    "sur les ressources du musée.")
para(doc,
    "Conduite au niveau intégrateur selon une démarche de recherche-action, l'étude a combiné "
    "l'observation directe du site, l'analyse documentaire du fonds et des entretiens avec le "
    "personnel de la Fondation. Les variables retenues — documentation des collections, "
    "médiation, accessibilité du public et valorisation économique — ont été mesurées par des "
    "indicateurs observables sur le terrain.")
para(doc,
    "[À COMPLÉTER APRÈS COLLECTE : synthèse des résultats obtenus sur le terrain — état de la "
    "documentation, volume des collections, pratiques de médiation observées.] Le diagnostic "
    "établi a conduit à proposer une intervention : MUSÉA, une plateforme numérique articulant "
    "un progiciel de gestion des collections, un site public de visite immersive en trois "
    "dimensions et en réalité augmentée, une restitution des généalogies reliant chaque objet "
    "à son souverain, un guide vocal fondé sur l'intelligence artificielle et un dispositif "
    "d'accès payant. La faisabilité de cette proposition a été établie par la réalisation d'un "
    "prototype fonctionnel, dont l'architecture et les mécanismes de sécurité sont documentés "
    "dans ce rapport.")

para(doc, "Mots clés : patrimoine culturel, médiation numérique, chefferies, musée, "
          "intelligence artificielle.", gras=True)

# ---- ABSTRACT
titre_liminaire(doc, "Abstract")
para(doc,
    "The heritage of the chiefdoms of Western Cameroon constitutes a major cultural corpus: "
    "artworks, regalia, architecture and dynastic lineages form an inseparable whole, in which "
    "each object draws its meaning from the history of the sovereign to whom it is attached. "
    "Mostly preserved in chiefdom museums, this heritage nevertheless remains hardly "
    "accessible: its documentation is still largely analogue, its mediation depends on the "
    "physical presence of a visitor and a guide, and the link between objects and the "
    "genealogies that carry them is rarely restored. The Maison Foudjem, a museum of the Jean "
    "Félicien Gacha Foundation, illustrates this situation.")
para(doc,
    "This study addresses the digital mediation of that heritage. It seeks to understand why "
    "the collections of the Maison Foudjem remain poorly accessible and poorly valorised, "
    "despite their richness. The general question investigates the causes of this situation; "
    "the specific questions concern current documentation and mediation practices, the state "
    "of the digitisation of objects and genealogies, and the effects of the absence of a "
    "digital device on audience and on the museum's resources.")
para(doc,
    "Conducted at the integrative level through an action-research approach, the study combined "
    "direct observation of the site, documentary analysis of the collection and interviews with "
    "the Foundation's staff. The selected variables — documentation of collections, mediation, "
    "public accessibility and economic valorisation — were measured through indicators "
    "observable in the field.")
para(doc,
    "[TO BE COMPLETED AFTER DATA COLLECTION: summary of field results.] The resulting diagnosis "
    "led to a proposed intervention: MUSÉA, a digital platform combining a collection "
    "management system, a public website offering an immersive visit in three dimensions and "
    "augmented reality, a rendering of the genealogies linking each object to its sovereign, an "
    "artificial-intelligence-based voice guide and a paid access mechanism. The feasibility of "
    "this proposal was established through a functional prototype, whose architecture and "
    "security mechanisms are documented in this report.")
para(doc, "Keywords: cultural heritage, digital mediation, chiefdoms, museum, artificial "
          "intelligence.", gras=True)

# ---- INDEX DES TABLEAUX / GRAPHIQUES / SOMMAIRE (champs a mettre a jour dans Word)
titre_liminaire(doc, "Index des tableaux")
para(doc, "[Dans Word : Références ▸ Insérer une table des illustrations ▸ Légende « Tableau ».]",
     italique=True, taille=11)

titre_liminaire(doc, "Index des graphiques")
para(doc, "[Dans Word : Références ▸ Insérer une table des illustrations ▸ Légende « Figure ».]",
     italique=True, taille=11)

titre_liminaire(doc, "Sommaire")
para(doc, "[Dans Word : Références ▸ Table des matières ▸ Table automatique (2 pages maximum).]",
     italique=True, taille=11)

# ================================================================ SECTION 3
# Corps : chiffres arabes, redemarrage a 1
s3 = nouvelle_section(doc)
format_pagination(s3, fmt="decimal", depart=1)
numeroter_bas_droite(s3)

# ---------------------------------------------------------------- INTRODUCTION
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
r = p.add_run("INTRODUCTION GÉNÉRALE")
r.bold = True
r.font.size = Pt(14)
r.font.name = POLICE

# --- 1. Contexte general
titre_2(doc, "1. Contexte général de l'étude")
para(doc,
    "Le patrimoine culturel constitue, pour les sociétés qui en héritent, à la fois une "
    "mémoire et une ressource. L'Organisation des Nations unies pour l'éducation, la science "
    "et la culture rappelle que ce patrimoine ne se limite pas aux monuments et aux "
    "collections d'objets : il comprend également les traditions, les savoir-faire et les "
    "expressions vivantes héritées des ancêtres et transmises aux descendants (UNESCO, 2003). "
    "Cette conception élargie s'applique particulièrement aux sociétés d'Afrique centrale, où "
    "l'objet et le récit qui le porte sont indissociables.")
para(doc,
    "Au Cameroun, les chefferies traditionnelles de l'Ouest occupent à cet égard une place "
    "singulière. Elles conservent un ensemble d'objets — trônes, masques, statues, insignes du "
    "pouvoir, textiles, architectures — dont la valeur ne tient pas seulement à la facture, "
    "mais au lien qui les rattache à une lignée. Un siège n'est pas un siège : il est celui "
    "d'un souverain déterminé, dont le règne, les alliances et la descendance donnent son sens "
    "à l'objet. La généalogie n'est donc pas un supplément documentaire ; elle est la clé de "
    "lecture du fonds.")
para(doc,
    "Ce patrimoine est aujourd'hui conservé dans des musées de chefferie et dans des "
    "fondations privées. La Fondation Jean Félicien Gacha s'inscrit dans cette démarche : "
    "elle œuvre à la conservation et à la transmission du patrimoine des chefferies "
    "camerounaises, notamment à travers la Maison Foudjem, où sont réunies des collections "
    "d'objets et une documentation sur les lignées auxquelles ils se rattachent.")
para(doc,
    "Dans le même temps, la médiation muséale connaît une transformation profonde. Le Conseil "
    "international des musées assigne désormais au musée une mission explicite de "
    "communication et de partage avec le public, au-delà de la seule conservation (ICOM, "
    "2022). Les technologies numériques — numérisation tridimensionnelle, réalité augmentée, "
    "restitution en ligne, médiation assistée par intelligence artificielle — ont "
    "considérablement élargi le champ des possibles. Elles permettent de donner accès à des "
    "collections sans déplacement, de restituer un objet dans son contexte et d'adresser un "
    "discours adapté à chaque visiteur.")
para(doc,
    "Cette évolution rencontre au Cameroun un contexte favorable sur le plan des usages. "
    "L'Union internationale des télécommunications observe une progression continue de la "
    "connectivité et de l'équipement mobile en Afrique subsaharienne (UIT, 2023), ce qui rend "
    "techniquement envisageable une médiation numérique adressée aussi bien au public local "
    "qu'à la diaspora.")
para(doc,
    "Un décalage subsiste néanmoins. Alors que les moyens techniques existent et que la "
    "demande de reconnexion à la mémoire des chefferies est réelle, les collections de "
    "nombreux musées de chefferie demeurent connues d'un cercle restreint. Leur documentation "
    "reste souvent manuscrite ou dispersée, leur médiation dépend de la présence simultanée "
    "d'un visiteur et d'un guide, et le lien entre l'objet et la lignée — qui en constitue "
    "pourtant l'intérêt majeur — n'est pas restitué au public. La Maison Foudjem n'échappe pas "
    "à ce constat. C'est ce décalage entre un fonds d'une grande richesse et une audience "
    "limitée qui fonde la problématique de la présente étude.")

# --- 2. Problematique
titre_2(doc, "2. Problématique de l'étude")

titre_3(doc, "2.1. Présentation du problème")
para(doc,
    "L'observation conduite à la Maison Foudjem pendant la période de stage fait apparaître "
    "une situation contrastée. D'un côté, le musée conserve un fonds d'objets rattachés aux "
    "lignées des chefferies, accompagné d'une connaissance orale et documentaire détenue par "
    "le personnel et par les dépositaires de la tradition. De l'autre, ce fonds ne produit "
    "qu'une audience restreinte et une contribution économique limitée.")
para(doc,
    "Trois faits empiriques ressortent de cette observation. Premièrement, la documentation "
    "des collections repose largement sur des supports analogiques et sur la mémoire des "
    "personnes : il n'existe pas de référentiel numérique unique décrivant chaque objet, sa "
    "localisation, son histoire et le souverain auquel il se rattache. Deuxièmement, la "
    "médiation est entièrement conditionnée par la coprésence physique : hors de la visite "
    "guidée, le patrimoine n'est ni consultable ni racontable, ce qui exclut de fait le public "
    "éloigné et la diaspora. Troisièmement, le lien entre l'objet et la généalogie, qui "
    "constitue la valeur d'usage principale du fonds, n'est restitué qu'oralement et de "
    "manière partielle.")
para(doc,
    "Ces constats produisent un effet cumulatif : le patrimoine est conservé, mais il est peu "
    "vu, peu documenté durablement et peu générateur de ressources pour la structure qui en "
    "assure la garde. La question qui se pose n'est donc pas celle de la richesse du fonds, "
    "mais celle des causes de sa faible diffusion.")

titre_3(doc, "2.2. Formulation du problème")
para(doc, "Question générale de l'étude", gras=True)
puce(doc, "Pourquoi le patrimoine conservé à la Maison Foudjem demeure-t-il faiblement "
          "accessible au public et faiblement valorisé, malgré la richesse de ses collections "
          "et des lignées qui s'y rattachent ?")
para(doc, "Questions spécifiques de l'étude", gras=True)
puce(doc, "Quels sont les dispositifs de documentation et de médiation actuellement mis en "
          "œuvre à la Maison Foudjem, et quels résultats produisent-ils ?")
puce(doc, "Quel est l'état de la documentation numérique des objets et des généalogies de la "
          "Maison Foudjem ?")
puce(doc, "En quoi l'absence d'un dispositif numérique de médiation limite-t-elle l'accès du "
          "public au patrimoine et la valorisation économique du fonds ?")

# --- 3. Hypotheses
titre_2(doc, "3. Hypothèses de l'étude")

titre_3(doc, "3.1. Hypothèse générale")
para(doc,
    "À la Maison Foudjem, le patrimoine demeure faiblement accessible et faiblement valorisé "
    "parce que la documentation des collections repose sur des supports analogiques et sur la "
    "mémoire des personnes, et parce que la médiation dépend de la coprésence physique d'un "
    "visiteur et d'un guide ; l'absence d'un système d'information restituant les objets et "
    "leurs généalogies prive le fonds d'une audience à distance et le musée d'une ressource "
    "associée à cette audience.")

titre_3(doc, "3.2. Hypothèses spécifiques")
puce(doc, "H1 : Les dispositifs de documentation et de médiation actuellement en usage à la "
          "Maison Foudjem sont fondés sur des supports analogiques et sur la transmission "
          "orale, ce qui limite leur portée au public physiquement présent.")
puce(doc, "H2 : La documentation numérique des objets et des généalogies de la Maison Foudjem "
          "est partielle et non structurée, ce qui empêche toute restitution en ligne "
          "cohérente du lien entre un objet et son souverain.")
puce(doc, "H3 : L'absence d'un dispositif numérique de médiation réduit l'audience du fonds au "
          "public local et prive la Fondation d'une source de revenus liée à la consultation "
          "à distance du patrimoine.")

# --- 4. Objectifs
titre_2(doc, "4. Objectifs de l'étude")

titre_3(doc, "4.1. Objectif général")
puce(doc, "Connaître les causes de la faible accessibilité et de la faible valorisation du "
          "patrimoine conservé à la Maison Foudjem.")

titre_3(doc, "4.2. Objectifs spécifiques")
puce(doc, "Connaître les dispositifs de documentation et de médiation mis en œuvre à la Maison "
          "Foudjem et en évaluer les résultats.")
puce(doc, "Évaluer l'état de la documentation numérique des objets et des généalogies de la "
          "Maison Foudjem.")
puce(doc, "Analyser le lien entre l'absence d'un dispositif numérique de médiation, l'audience "
          "du fonds et sa valorisation économique.")

# --- 5. Justification
titre_2(doc, "5. Justification de l'étude")

titre_3(doc, "5.1. Au plan scientifique")
para(doc,
    "La médiation numérique du patrimoine est un champ documenté, mais les travaux portent "
    "majoritairement sur les grandes institutions muséales occidentales. Les musées de "
    "chefferie d'Afrique centrale, dont le fonctionnement repose sur l'articulation entre "
    "l'objet et la lignée, restent peu étudiés sous cet angle. Cette étude apporte une "
    "connaissance située sur les causes de la faible diffusion d'un fonds patrimonial dans ce "
    "contexte précis. Elle documente en outre une articulation peu traitée dans la "
    "littérature : celle qui relie la collection muséale à la structure généalogique qui lui "
    "donne son sens, et qui conditionne la conception de tout dispositif de médiation adapté à "
    "ces musées.")

titre_3(doc, "5.2. Au plan pratique")
para(doc,
    "Sur le plan pratique, l'étude fournit à la Fondation Jean Félicien Gacha un diagnostic "
    "des causes de la faible audience de la Maison Foudjem, appuyé sur des observations de "
    "terrain. Elle débouche sur une proposition d'intervention dont la faisabilité est "
    "établie par un prototype fonctionnel. Les enseignements tirés sont transposables aux "
    "autres musées de chefferie confrontés à la même situation, ainsi qu'aux acteurs publics "
    "et privés engagés dans la valorisation du patrimoine culturel camerounais.")

# --- 6. Delimitation
titre_2(doc, "6. Délimitation de l'étude")

titre_3(doc, "6.1. Délimitation géographique")
para(doc,
    "L'étude est circonscrite à la Maison Foudjem, musée de la Fondation Jean Félicien Gacha, "
    "au Cameroun. Elle ne prétend pas couvrir l'ensemble des musées de chefferie du pays : les "
    "conclusions valent pour ce site, et leur transposition à d'autres structures suppose des "
    "vérifications préalables.")

titre_3(doc, "6.2. Délimitation temporelle")
para(doc,
    "Les données ont été collectées durant la période de stage effectuée à la Maison Foudjem "
    "au cours de l'année académique 2025-2026. [Préciser les dates exactes de début et de fin "
    "du stage.] L'analyse porte sur la situation observée pendant cette période.")

titre_3(doc, "6.3. Délimitation thématique et conceptuelle")
para(doc,
    "L'étude porte sur la médiation numérique du patrimoine, entendue comme l'ensemble des "
    "dispositifs techniques permettant de documenter, de restituer et de raconter un fonds "
    "patrimonial à un public, sur un support numérique. Elle retient quatre dimensions : la "
    "documentation des collections, la médiation, l'accessibilité du public et la valorisation "
    "économique. Sont écartées les questions de conservation matérielle des objets, de "
    "restauration, de statut juridique des biens et de politique culturelle nationale, qui "
    "relèvent d'autres disciplines.")

titre_3(doc, "6.4. Délimitation de la population")
para(doc,
    "La population de l'étude est constituée du fonds patrimonial de la Maison Foudjem — "
    "objets et lignées documentées — ainsi que du personnel de la Fondation intervenant dans "
    "sa conservation et sa médiation. Les visiteurs occasionnels n'ont pas été enquêtés de "
    "manière systématique, ce qui limite la portée des conclusions relatives à la réception "
    "du public.")

# --- 7. Plan
titre_2(doc, "7. Plan du rapport")
para(doc,
    "Ce rapport est organisé en quatre chapitres. Le premier expose le cadre conceptuel et "
    "l'état de l'art : il clarifie les notions de patrimoine culturel, de médiation numérique, "
    "de généalogie et de valorisation, et situe l'étude au regard des travaux existants. Le "
    "deuxième présente la méthodologie retenue : nature de l'étude, variables et indicateurs, "
    "population, échantillonnage et outils de collecte. Le troisième présente le site de "
    "l'étude — la Maison Foudjem — ainsi que les données collectées et les résultats obtenus, "
    "y compris l'implémentation réalisée. Le quatrième établit le diagnostic de la situation, "
    "vérifie les hypothèses et présente l'intervention proposée ainsi que sa justification. "
    "Une conclusion générale rappelle les apports du travail et ouvre des perspectives.")

# ---------------------------------------------------------------- chapitres
from chapitres_1_2 import chapitre1, chapitre2
from chapitres_3_4 import chapitre3, chapitre4, conclusion_generale, references, annexes

chapitre1(doc)
chapitre2(doc)
chapitre3(doc)
chapitre4(doc)
conclusion_generale(doc)
references(doc)
annexes(doc)

# ---------------------------------------------------------------- sauvegarde
sortie = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                      "Rapport_DASSI_KAMGANG_Max_Brian.docx")
doc.save(sortie)
print("Document genere :", sortie)
